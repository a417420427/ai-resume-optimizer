import io
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════════════
#  Resume Style Templates
# ═══════════════════════════════════════════════════════════════════════

TEMPLATES = {
    "modern": {
        "name": "现代简约",
        "primary":   (0x18, 0x97, 0xFF),   # 主色蓝
        "secondary": (0x1A, 0x1A, 0x2E),   # 深蓝黑
        "accent":    (0x52, 0xC4, 0x1A),   # 绿色点缀
        "text":      (0x33, 0x33, 0x33),   # 正文
        "muted":     (0x99, 0x99, 0x99),   # 辅助灰
        "divider":   (0x18, 0x97, 0xFF),   # 分割线
        "bg_light":  (0xF5, 0xFA, 0xFF),   # 浅蓝背景
        "name_size": 22,
        "section_size": 13,
        "body_size": 10.5,
        "contact_size": 9,
    },
    "classic": {
        "name": "经典商务",
        "primary":   (0x2C, 0x3E, 0x50),
        "secondary": (0x1A, 0x1A, 0x2E),
        "accent":    (0xE7, 0x4C, 0x3C),
        "text":      (0x33, 0x33, 0x33),
        "muted":     (0x99, 0x99, 0x99),
        "divider":   (0x2C, 0x3E, 0x50),
        "bg_light":  (0xF8, 0xF8, 0xF8),
        "name_size": 22,
        "section_size": 13,
        "body_size": 10.5,
        "contact_size": 9,
    },
    "minimal": {
        "name": "极简风格",
        "primary":   (0x55, 0x55, 0x55),
        "secondary": (0x33, 0x33, 0x33),
        "accent":    (0x00, 0x00, 0x00),
        "text":      (0x44, 0x44, 0x44),
        "muted":     (0xAA, 0xAA, 0xAA),
        "divider":   (0xCC, 0xCC, 0xCC),
        "bg_light":  (0xFA, 0xFA, 0xFA),
        "name_size": 20,
        "section_size": 12,
        "body_size": 10,
        "contact_size": 8.5,
    },
}


class DocService:

    # ── DOCX ──────────────────────────────────────────────────────────

    @classmethod
    def to_docx(cls, markdown_text: str, template: str = "modern") -> tuple[bytes, str]:
        """Generate a properly styled DOCX resume."""
        t = TEMPLATES.get(template, TEMPLATES["modern"])
        doc = Document()

        # ── Page Setup ──
        section = doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # ── Default paragraph font ──
        style = doc.styles['Normal']
        font = style.font
        font.name = 'PingFang SC'
        font.size = Pt(t["body_size"])
        font.color.rgb = cls._rgb(*t["text"])
        pf = style.paragraph_format
        pf.space_after = Pt(1)
        pf.line_spacing = 1.3
        # East-Asian font
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'PingFang SC')

        def add_text(para, text, bold=False, italic=False, size=None, color=None):
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = cls._rgb(*color)
            run.font.name = 'PingFang SC'
            rPr = run._element.get_or_add_rPr()
            rF = rPr.find(qn('w:rFonts'))
            if rF is None:
                rF = OxmlElement('w:rFonts')
                rPr.append(rF)
            rF.set(qn('w:eastAsia'), 'PingFang SC')
            return run

        def set_para_spacing(para, before=0, after=0):
            pf = para.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)

        def add_divider(para):
            """Add a thin colored line below the paragraph."""
            pPr = para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '2')
            c = t["divider"]
            bottom.set(qn('w:color'), f'{c[0]:02X}{c[1]:02X}{c[2]:02X}')
            pBdr.append(bottom)
            pPr.append(pBdr)

        def extract_bold(text):
            """Split text into [(bold, segment), ...]"""
            parts = re.split(r'(\*\*.*?\*\*)', text)
            result = []
            for p in parts:
                if p.startswith('**') and p.endswith('**'):
                    result.append((True, p[2:-2]))
                elif p:
                    result.append((False, p))
            return result

        def write_inline(para, text, size=None, color=None):
            """Write text with **bold** support."""
            if not text:
                return
            for bold, seg in extract_bold(text):
                add_text(para, seg, bold=bold, size=size or t["body_size"], color=color)

        # ── Parse Markdown ──
        lines = markdown_text.strip().split('\n')
        i = 0
        while i < len(lines):
            raw = lines[i]
            line = raw.strip()
            i += 1
            if not line:
                continue

            # ── Name (H1) ──
            if line.startswith('# ') and not line.startswith('## '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_spacing(p, 18, 2)
                add_text(p, line[2:], bold=True, size=t["name_size"], color=t["secondary"])

            # ── Contact info ──
            elif cls._is_contact(line):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_spacing(p, 2, 10)
                add_text(p, line, size=t["contact_size"], color=t["muted"])

            # ── Section header (H2) ──
            elif line.startswith('## '):
                p = doc.add_paragraph()
                set_para_spacing(p, 14, 4)
                add_text(p, line[3:], bold=True, size=t["section_size"], color=t["primary"])
                add_divider(p)

            # ── Sub-section (H3) ──
            elif line.startswith('### '):
                p = doc.add_paragraph()
                set_para_spacing(p, 8, 2)
                add_text(p, line[4:], bold=True, size=11, color=(0x44, 0x44, 0x44))

            # ── Bullet ──
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.first_line_indent = Cm(-0.35)
                set_para_spacing(p, 1, 1)
                write_inline(p, '• ' + line[2:])

            # ── Numbered ──
            elif re.match(r'^\d+[\.\)]\s', line):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                set_para_spacing(p, 1, 1)
                write_inline(p, line)

            # ── Regular paragraph ──
            else:
                p = doc.add_paragraph()
                set_para_spacing(p, 1, 1)
                write_inline(p, line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue(), "optimized-resume.docx"

    # ── PDF ──────────────────────────────────────────────────────────

    @classmethod
    def to_pdf(cls, markdown_text: str, template: str = "modern") -> tuple[bytes, str]:
        """Generate a proper 2-column layout resume PDF."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm, cm
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Frame,
            PageTemplate, BaseDocTemplate,
        )
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        t = TEMPLATES.get(template, TEMPLATES["modern"])

        # Font
        font_path = cls._ensure_cn_font()
        pdfmetrics.registerFont(TTFont('CN', font_path))
        pdfmetrics.registerFont(TTFont('CNB', font_path))

        def c(r, g, b):
            return colors.Color(r / 255, g / 255, b / 255)

        pri_c = c(*t["primary"])
        sec_c = c(*t["secondary"])
        txt_c = c(*t["text"])
        mut_c = c(*t["muted"])

        # ── Styles ──
        S = {}

        S['name'] = ParagraphStyle('N', fontName='CNB', fontSize=t["name_size"],
                                   leading=t["name_size"] * 1.2, alignment=1,
                                   spaceAfter=2*mm, textColor=sec_c)

        S['contact'] = ParagraphStyle('C', fontName='CN', fontSize=t["contact_size"],
                                      leading=t["contact_size"] * 1.5, alignment=1,
                                      spaceAfter=6*mm, textColor=mut_c)

        S['section'] = ParagraphStyle('Sec', fontName='CNB', fontSize=t["section_size"],
                                      leading=t["section_size"] * 1.4,
                                      spaceBefore=5*mm, spaceAfter=1*mm,
                                      textColor=pri_c)

        S['subsection'] = ParagraphStyle('Sub', fontName='CNB', fontSize=10.5,
                                         leading=15, spaceBefore=3*mm,
                                         spaceAfter=1*mm, textColor=c(0x44, 0x44, 0x44))

        S['body'] = ParagraphStyle('B', fontName='CN', fontSize=t["body_size"],
                                   leading=t["body_size"] * 1.65, spaceAfter=1*mm,
                                   textColor=txt_c)

        S['bullet'] = ParagraphStyle('BL', fontName='CN', fontSize=t["body_size"],
                                     leading=t["body_size"] * 1.65, spaceAfter=0.5*mm,
                                     leftIndent=14, bulletIndent=0, textColor=txt_c)

        def esc(text):
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            return text

        # ── Build story ──
        story = []
        lines = markdown_text.strip().split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            i += 1
            if not line:
                continue

            try:
                if line.startswith('# ') and not line.startswith('## '):
                    story.append(Paragraph(esc(line[2:]), S['name']))

                elif cls._is_contact(line):
                    story.append(Paragraph(esc(line), S['contact']))

                elif line.startswith('## '):
                    story.append(Paragraph(esc(line[3:]), S['section']))
                    # Divider table
                    div_table = Table(
                        [['']],
                        colWidths=[170*mm],
                        style=TableStyle([
                            ('LINEBELOW', (0, 0), (-1, -1), 0.8, pri_c),
                            ('TOPPADDING', (0, 0), (-1, -1), 0),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                        ])
                    )
                    story.append(div_table)

                elif line.startswith('### '):
                    story.append(Paragraph(esc(line[4:]), S['subsection']))

                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f'<bullet>&bull;</bullet>{esc(line[2:])}', S['bullet']))

                elif re.match(r'^\d+[\.\)]\s', line):
                    text = re.sub(r'^\d+[\.\)]\s', '', line)
                    story.append(Paragraph(f'<bullet>&bull;</bullet>{esc(text)}', S['bullet']))

                else:
                    story.append(Paragraph(esc(line), S['body']))

            except Exception:
                story.append(Paragraph(esc(line), S['body']))

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            topMargin=20*mm, bottomMargin=18*mm,
            leftMargin=25*mm, rightMargin=25*mm,
        )
        doc.build(story)
        buf.seek(0)
        return buf.getvalue(), "optimized-resume.pdf"

    # ── Utilities ─────────────────────────────────────────────────────

    @staticmethod
    def _is_contact(line: str) -> bool:
        return bool(
            '@' in line
            or '电话' in line
            or '邮箱' in line
            or re.match(r'^1[3-9]\d', line)
            or '|' in line and re.search(r'\d', line)
            and len(line) < 60
        )

    @staticmethod
    def _rgb(r, g, b) -> RGBColor:
        return RGBColor(r, g, b)

    @classmethod
    def get_templates(cls) -> list[dict]:
        return [{"key": k, "name": v["name"]} for k, v in TEMPLATES.items()]

    @classmethod
    def _ensure_cn_font(cls) -> str:
        ttf_path = '/tmp/HeitiSC.ttf'
        if not os.path.exists(ttf_path):
            from fontTools.ttLib import TTCollection
            ttc = TTCollection('/System/Library/Fonts/STHeiti Medium.ttc')
            ttc[1].save(ttf_path)
        return ttf_path
